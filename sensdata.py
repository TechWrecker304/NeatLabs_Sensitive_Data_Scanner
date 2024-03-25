import os
import re
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import ttkbootstrap as ttk
from docx import Document
from openpyxl import load_workbook

def scan_directory(directory, patterns):
    results = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            try:
                content = ""
                if file.endswith('.txt'):
                    with open(file_path, 'r') as f:
                        content = f.read()
                elif file.endswith('.docx'):
                    doc = Document(file_path)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                elif file.endswith('.xlsx'):
                    wb = load_workbook(file_path)
                    ws = wb.active
                    content = '\n'.join(['\t'.join([str(cell.value) for cell in row]) for row in ws.iter_rows()])

                for pattern, description in patterns.items():
                    matches = re.findall(pattern, content)
                    if matches:
                        results.append((file_path, description, matches))
            except:
                pass
    return results

def copy_files(files, destination):
    for file_path in files:
        shutil.copy2(file_path, destination)

def generate_report(results, report_path):
    doc = Document()
    doc.add_heading('Sensitive Data Detection Report', 0)

    for file_path, description, matches in results:
        doc.add_heading(file_path, 1)
        doc.add_paragraph(f"Sensitive Data Type: {description}")
        doc.add_paragraph(f"Matches: {', '.join(matches)}")
        doc.add_page_break()

    doc.save(report_path)

def scan_files():
    directory = filedialog.askdirectory(title="Select Directory to Scan")
    if directory:
        selected_patterns = {pattern: description for pattern, description in data_patterns.items() if pattern_vars[pattern].get()}
        results = scan_directory(directory, selected_patterns)
        result_text.delete('1.0', tk.END)
        if results:
            for file_path, description, matches in results:
                result_text.insert(tk.END, f"File: {file_path}\nSensitive Data: {description}\nMatches: {matches}\n\n")
        else:
            result_text.insert(tk.END, "No sensitive data found.")
        messagebox.showinfo("Scan Complete", "Scanning completed successfully.")

def copy_selected_files():
    destination = filedialog.askdirectory(title="Select Destination Directory")
    if destination:
        selected_files = [file_info.split('\n')[0].split(': ')[1] for file_info in result_text.get('1.0', tk.END).strip().split('\n\n') if file_info.strip()]
        copy_files(selected_files, destination)
        messagebox.showinfo("Copy Complete", "Selected files have been copied.")

def clear_results():
    result_text.delete('1.0', tk.END)

def generate_report_action():
    report_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if report_path:
        results = [(file_info.split('\n')[0].split(': ')[1], file_info.split('\n')[1].split(': ')[1], file_info.split('\n')[2].split(': ')[1].strip('[]').split(', ')) for file_info in result_text.get('1.0', tk.END).strip().split('\n\n') if file_info.strip()]
        generate_report(results, report_path)
        messagebox.showinfo("Report Generated", f"Report has been generated and saved as {report_path}")

# Create the main window
window = ttk.Window(themename="darkly")
window.title("NeatLabs Sensitive Data Detection Scanner")
window.geometry("1200x1200")

# Create the data patterns frame
pattern_frame = ttk.LabelFrame(window, text="Sensitive Data Patterns", padding=10)
pattern_frame.pack(pady=10)

data_patterns = {
    r'\b\d{3}-\d{2}-\d{4}\b': 'Social Security Number',
    r'\b(?:\d{4}[-\s]?){3}\d{4}\b': 'Credit Card Number',
    r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b': 'Email Address',
    r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b': 'IP Address',
    r'\b(?:\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}\b': 'Phone Number',
    r'\b[A-Za-z]{2}\d{2}[A-Za-z]{2}\d{4}\b': 'Passport Number'
}

pattern_vars = {}
for pattern, description in data_patterns.items():
    var = tk.BooleanVar(value=True)
    pattern_vars[pattern] = var
    checkbox = ttk.Checkbutton(pattern_frame, text=description, variable=var)
    checkbox.pack(anchor=tk.W)

# Create the scan button
scan_button = ttk.Button(window, text="Scan Files", command=scan_files)
scan_button.pack(pady=10)

# Create the result text area
result_frame = ttk.LabelFrame(window, text="Scan Results", padding=10)
result_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

result_text = tk.Text(result_frame, wrap=tk.WORD)
result_text.pack(fill=tk.BOTH, expand=True)

# Create the action buttons
action_frame = ttk.Frame(window)
action_frame.pack(pady=10)

copy_button = ttk.Button(action_frame, text="Copy Selected Files", command=copy_selected_files)
copy_button.pack(side=tk.LEFT, padx=5)

clear_button = ttk.Button(action_frame, text="Clear Results", command=clear_results)
clear_button.pack(side=tk.LEFT, padx=5)

report_button = ttk.Button(action_frame, text="Generate Report", command=generate_report_action)
report_button.pack(side=tk.LEFT, padx=5)

# Start the main event loop
window.mainloop()